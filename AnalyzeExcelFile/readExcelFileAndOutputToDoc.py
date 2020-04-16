import xlrd
import docx
import os
from docx.shared import Inches,Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 新建doc文档
doc = docx.Document()

# 将正文英文字体设置成
doc.styles['Normal'].font.name = u'Times New Roman'
# 将正文中文字体设置成宋体
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

# 获取正文样式
style = doc.styles['Normal']
# 获取段落样式
paragraph_format = style.paragraph_format
# 首行缩进0.74厘米，即2个字符
paragraph_format.first_line_indent = Cm(0.74)

# 设置1级标题的字体为Times New Roman， 且无缩进
style_heading1 = doc.styles['Heading 1']
heading1_format = style_heading1.paragraph_format
heading1_format.first_line_indent = Cm(0)
style_heading1.font.name = 'Times New Roman'

# 设置2级标题的字体为Times New Roman， 且无缩进
style_heading2 = doc.styles['Heading 2']
heading2_format = style_heading2.paragraph_format
heading2_format.first_line_indent = Cm(0)
style_heading2.font.name = 'Times New Roman'

# open the .xlsx file
book = xlrd.open_workbook('resource/helk dashboard visualization_补充完整.xlsx')

sheet1 = book.sheets()[0]


dashboard = ""

for rowNum in range(1, sheet1.nrows):

    curr_dashboard = sheet1.cell(rowNum, 0).value
    if(curr_dashboard is not ""):
        dashboard = curr_dashboard
        doc.add_heading(dashboard, 1)
    
    visualization = sheet1.cell(rowNum, 1).value
    visualization_type = sheet1.cell(rowNum, 2).value
    description = sheet1.cell(rowNum, 3).value

    if(visualization is not ""):
        doc.add_heading(visualization, 2)
        
        pic_path = 'resource/pic/' + visualization + '.png'
        if os.path.exists(pic_path):
            
            doc.add_paragraph('该visualization的类型为' + visualization_type + ',其作用是' + description + '。其界面如下图所示：', style='Body Text')
            # 指定图片宽度为 4英寸, 图片高度会自动调整
            doc.add_picture(pic_path, width=Inches(4))
            last_paragraph = doc.paragraphs[-1]
            #图片居中设置
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 输出已经分析的visualization名字到指定文件
            with open("resource/visualization.txt", "at", encoding="utf-8") as file:
                file.write(visualization+'\n')
        else:
            doc.add_paragraph('该visualization的类型为' + visualization_type + ',其作用是' + description, style='Body Text')
            print(pic_path + "不存在")


doc.save('resource/output.docx')


