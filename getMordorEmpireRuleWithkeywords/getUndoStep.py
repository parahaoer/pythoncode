import xlrd
import docx

# open the .xlsx file
book = xlrd.open_workbook(u'resource/apt3_mordor_playbook(2).xlsx')
sheet2 = book.sheets()[0]
stepList = sheet2.col_values(0)
stepList = [item.strip() for item in stepList]

# print(undoList)


undoList = []
for line in open('resource/undo.txt', 'r', encoding='gb18030'):
    undoList.append(line.strip())


undoRowNumList = [count for count, item in enumerate(stepList) if item in undoList]

print(undoRowNumList)


# 新建doc文档
doc = docx.Document()

for id, rowNum in enumerate(undoRowNumList):
    
    evalStep = undoList[id]
    technique_code = sheet2.cell(rowNum, 3).value
    technique = sheet2.cell(rowNum, 4).value
    tactic = sheet2.cell(rowNum, 2).value
    doc.add_paragraph(u'Eval Step：' + evalStep)
    doc.add_paragraph(u'检测目标：' + technique_code + u' - ' + technique)
    doc.add_paragraph(u'所属tactics：TA0010 – ' + tactic)
    doc.add_paragraph(u"检索日志来源：PowerShell")
    doc.add_paragraph(u"检索规则")
    doc.add_paragraph(u"\n")

doc.save('resource/output.docx')

