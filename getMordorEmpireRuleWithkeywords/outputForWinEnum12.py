from resource.code.winEnum_12_E_12_info import *
import docx

# 新建doc文档
doc = docx.Document()

for id, keyword in enumerate(rule_3_list):
    eval_step = eval_step_list[id]
    technique_id = technique_id_list[id]
    technique = technique_list[id]
    tactic = tactic_list[id]
    doc.add_paragraph(u'Eval Step：' + eval_step)
    doc.add_paragraph(u'检测目标：' + technique_id + u' - ' + technique)
    doc.add_paragraph(u'所属tactics：TA0007 – ' + tactic)
    doc.add_paragraph(u"检索日志来源：PowerShell")
    doc.add_paragraph(u"检索规则")
    doc.add_paragraph(u"\n")
    doc.add_paragraph(keyword)
    doc.add_paragraph(u"\n")
doc.save('resource/outputForWinEnum.docx')

