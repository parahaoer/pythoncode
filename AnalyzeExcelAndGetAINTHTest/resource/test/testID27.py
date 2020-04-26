import docx
from docx import Document
import re

path = 'resource/doc/25-31.docx'

doc = Document(path)

for id in range(1, 60):
    for j, paragraph in enumerate(doc.paragraphs):
        if id == 27 and re.match('^' + str(id) , paragraph.text):

            print(doc.paragraphs[j].text)
            continue