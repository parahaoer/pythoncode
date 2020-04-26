import xlrd
import docx
from docx import Document
import re
from xlrd import open_workbook
from xlutils.copy import copy
import os

def func(file_path):
    rb = open_workbook('resource/excel/getTestItem.xls')
    wb = copy(rb)
    sheet1 = rb.sheets()[0]

    doc = Document(file_path)

    # open the .xlsx file
    book = xlrd.open_workbook('resource/excel/对mordor small_dataset的查询 目录_增加.xlsx')

    sheet1 = book.sheets()[0]

    for id in range(1, sheet1.nrows):
        for j, paragraph in enumerate(doc.paragraphs):
            if re.match('^' + str(id) + u'、\\s*$' , paragraph.text) or re.match('^' + str(id) + '\.\\s*$' , paragraph.text):
                print(id)
                # print(doc.paragraphs[j+1].text)
  
                procedure = sheet1.cell(id, 5).value
                dataset = sheet1.cell(id, 1).value
                testItem = u'在数据集' + dataset + u'中， 检测procedure：' + procedure
                wb.get_sheet(0).write(id, 1, testItem)

                print(testItem)
    
                continue
    # rb = open_workbook打开那个excel文件，就保存到那个excel文件，否则写入到目标文件的内容不完整。
    wb.save('resource/excel/getTestItem.xls')
        




dir_path = 'resource/doc'

for file_name in os.listdir(dir_path):
    filePath = dir_path + '/' + file_name
    func(filePath)