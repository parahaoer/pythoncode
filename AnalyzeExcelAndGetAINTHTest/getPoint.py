import xlrd
import docx
from docx import Document
import re
from xlrd import open_workbook
from xlutils.copy import copy
import os

def func(file_path):
    rb = open_workbook('resource/excel/forOutput.xls')
    wb = copy(rb)
    sheet1 = rb.sheets()[0]

    doc = Document(file_path)

    # open the .xlsx file
    book = xlrd.open_workbook('resource/excel/对mordor small_dataset的查询 目录_增加.xlsx')

    sheet1 = book.sheets()[0]

    for id in range(1, sheet1.nrows):
        for j, paragraph in enumerate(doc.paragraphs):
            if re.match('^' + str(id) + u'、\\s*$' , paragraph.text) or re.match('^' + str(id) + '\.\\s*$' , paragraph.text):
                print(id)
                # print(doc.paragraphs[j+1].text)
                desc = sheet1.cell(id, 6).value
                tactic = sheet1.cell(id, 2).value
                technique = sheet1.cell(id, 4).value
                technique_id = sheet1.cell(id, 3).value
                procedure = sheet1.cell(id, 5).value
                
                point = doc.paragraphs[j+1].text + u'在数据集中检测行为：' + desc + u'。该行为对应ATT&CK中的' + tactic + u'下的'  + technique_id + ' ' + technique + u'下的procudure：' +  procedure
                wb.get_sheet(0).write(id, 4, point)

                print(point)
    
                continue
    # rb = open_workbook打开那个excel文件，就保存到那个excel文件，否则写入到目标文件的内容不完整。
    wb.save('resource/excel/forOutput.xls')
        




dir_path = 'resource/doc'

for file_name in os.listdir(dir_path):
    filePath = dir_path + '/' + file_name
    func(filePath)