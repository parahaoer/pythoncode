import docx
from docx import Document
import re

path = 'resource/doc/40-47.docx'

doc = Document(path)

for id in range(1, 60):
    for j, paragraph in enumerate(doc.paragraphs):
        if re.match('^' + str(id) + u'、\\s*$' , paragraph.text) or re.match('^' + str(id) + '\.\\s*$' , paragraph.text):
            print(id)

            print(doc.paragraphs[j+1].text)
            continue